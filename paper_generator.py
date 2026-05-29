from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE_KEYS = {
    "題目": "title",
    "英文題目": "title_en",
    "作者": "author",
    "單位": "department",
    "系所": "department",
    "類型": "document_type",
    "學位": "document_type",
    "學校": "organization",
    "機構": "organization",
    "年份": "year",
    "月份": "month",
}

SECTION_ALIASES = {
    "中文摘要": "abstract_zh",
    "摘要": "abstract_zh",
    "執行摘要": "abstract_zh",
    "英文摘要": "abstract_en",
    "Abstract": "abstract_en",
    "誌謝": "acknowledgements",
    "謝誌": "acknowledgements",
    "參考文獻": "references",
    "附錄": "appendices",
}

SPECIAL_HEADINGS = set(SECTION_ALIASES.keys())


@dataclass(frozen=True)
class TemplateSpec:
    key: str
    label: str
    document_type: str
    organization_default: str
    abstract_title: str
    english_abstract_title: str
    toc_title: str
    references_title: str
    appendix_title: str
    include_english_abstract: bool
    include_figure_lists: bool
    cover_hint: str


TEMPLATES: dict[str, TemplateSpec] = {
    "thesis": TemplateSpec(
        key="thesis",
        label="論文版",
        document_type="碩士論文",
        organization_default="國立嘉義大學",
        abstract_title="摘  要",
        english_abstract_title="Abstract",
        toc_title="目次",
        references_title="參考文獻",
        appendix_title="附錄",
        include_english_abstract=True,
        include_figure_lists=True,
        cover_hint="適合學術論文、研究計畫、正式研究報告。",
    ),
    "report": TemplateSpec(
        key="report",
        label="專題報告版",
        document_type="專題報告",
        organization_default="未命名單位",
        abstract_title="摘要",
        english_abstract_title="Abstract",
        toc_title="目錄",
        references_title="參考資料",
        appendix_title="附錄",
        include_english_abstract=False,
        include_figure_lists=True,
        cover_hint="適合課堂專題、結案報告、技術報告。",
    ),
    "proposal": TemplateSpec(
        key="proposal",
        label="商業提案版",
        document_type="商業提案",
        organization_default="未命名公司",
        abstract_title="執行摘要",
        english_abstract_title="Executive Summary",
        toc_title="提案目錄",
        references_title="參考資料",
        appendix_title="補充附件",
        include_english_abstract=False,
        include_figure_lists=False,
        cover_hint="適合提案書、募資簡報文稿、商業合作文件。",
    ),
}

TEMPLATE_OPTIONS = [
    {
        "key": spec.key,
        "label": spec.label,
        "documentType": spec.document_type,
        "hint": spec.cover_hint,
    }
    for spec in TEMPLATES.values()
]


@dataclass
class TableBlock:
    title: str
    headers: list[str]
    rows: list[list[str]]
    note: str = ""


@dataclass
class FigureBlock:
    title: str
    path: str
    note: str = ""
    width_cm: float = 14.0


@dataclass
class SectionBlock:
    title: str
    blocks: list[object] = field(default_factory=list)


@dataclass
class ChapterBlock:
    title: str
    sections: list[SectionBlock] = field(default_factory=list)


@dataclass
class PaperContent:
    metadata: dict[str, str]
    abstract_zh: str = ""
    keywords_zh: list[str] = field(default_factory=list)
    abstract_en: str = ""
    keywords_en: list[str] = field(default_factory=list)
    acknowledgements: str = ""
    chapters: list[ChapterBlock] = field(default_factory=list)
    references: list[str] = field(default_factory=list)
    appendices: list[tuple[str, str]] = field(default_factory=list)


class ParseError(ValueError):
    pass


def generate_professional_docx(
    raw_text: str,
    output_path: str | Path,
    template_key: str = "thesis",
) -> Path:
    template = TEMPLATES.get(template_key, TEMPLATES["thesis"])
    paper = parse_document_text(raw_text, template)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document)
    _define_styles(document)

    _build_cover(document, paper, template)

    front_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(front_section)
    _set_page_number_format(front_section, fmt="lowerRoman", start=1)
    _add_page_number(front_section)

    _build_abstract_zh(document, paper, template)
    if template.include_english_abstract:
        _build_abstract_en(document, paper, template)
    if paper.acknowledgements:
        _build_center_title(document, "誌謝")
        _write_multiline_paragraphs(document, paper.acknowledgements)
        document.add_page_break()

    _build_center_title(document, template.toc_title)
    _insert_toc_field(document.add_paragraph(), r'TOC \o "1-3" \h \z \u')
    document.add_page_break()

    if template.include_figure_lists:
        _build_center_title(document, "表次")
        _insert_toc_field(document.add_paragraph(), r'TOC \h \z \f t')
        document.add_page_break()

        _build_center_title(document, "圖次")
        _insert_toc_field(document.add_paragraph(), r'TOC \h \z \f f')
        document.add_page_break()

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(body_section)
    _set_page_number_format(body_section, fmt="decimal", start=1)
    _add_page_number(body_section)

    figure_counters: dict[int, int] = {}
    table_counters: dict[int, int] = {}

    for chapter_index, chapter in enumerate(paper.chapters, start=1):
        _add_chapter_heading(document, chapter.title, chapter_index > 1)
        if not chapter.sections:
            chapter.sections.append(SectionBlock(title="第一節 內容"))
        for section in chapter.sections:
            _add_section_heading(document, section.title)
            for block in section.blocks:
                if isinstance(block, str):
                    if block == "":
                        document.add_paragraph("")
                    else:
                        p = document.add_paragraph(style="BodyTextCustom")
                        p.add_run(block)
                elif isinstance(block, TableBlock):
                    caption = _next_caption("表", chapter_index, table_counters, block.title)
                    _add_caption(document, caption, toc_id="t")
                    _add_table(document, block)
                elif isinstance(block, FigureBlock):
                    caption = _next_caption("圖", chapter_index, figure_counters, block.title)
                    _add_caption(document, caption, toc_id="f")
                    _add_figure(document, block)

    _build_references(document, paper.references, template.references_title)
    _build_appendices(document, paper.appendices, template.appendix_title)
    _mark_fields_for_update(document)
    document.save(output)
    return output


def generate_paper_docx(raw_text: str, output_path: str | Path) -> Path:
    return generate_professional_docx(raw_text, output_path, template_key="thesis")


def parse_document_text(raw_text: str, template: TemplateSpec) -> PaperContent:
    lines = _normalize_raw_text(raw_text).split("\n")
    metadata: dict[str, str] = {
        "organization": template.organization_default,
        "document_type": template.document_type,
        "department": "",
        "author": "",
        "title": "未命名文件",
        "title_en": "Untitled Document",
        "year": "",
        "month": "",
    }

    chapters: list[ChapterBlock] = []
    references: list[str] = []
    appendices: list[tuple[str, str]] = []
    abstract_zh_lines: list[str] = []
    abstract_en_lines: list[str] = []
    acknowledgements_lines: list[str] = []
    keywords_zh: list[str] = []
    keywords_en: list[str] = []

    active_special: str | None = None
    current_chapter: ChapterBlock | None = None
    current_section: SectionBlock | None = None

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            if current_section and current_section.blocks:
                current_section.blocks.append("")
            i += 1
            continue

        meta_match = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if meta_match and meta_match.group(1).strip() in TITLE_KEYS:
            key = TITLE_KEYS[meta_match.group(1).strip()]
            metadata[key] = meta_match.group(2).strip()
            i += 1
            continue

        if line.startswith("關鍵詞："):
            keywords_zh = [x.strip() for x in line.split("：", 1)[1].split("、") if x.strip()]
            i += 1
            continue

        if line.startswith("Keywords:"):
            keywords_en = [x.strip() for x in line.split(":", 1)[1].split(",") if x.strip()]
            i += 1
            continue

        section_match = re.match(r"^#\s+(.+)$", line)
        if section_match:
            title = section_match.group(1).strip()
            active_special = SECTION_ALIASES.get(title)
            if active_special is None:
                current_chapter = ChapterBlock(title=title)
                chapters.append(current_chapter)
                current_section = None
            else:
                current_section = None
            i += 1
            continue

        sub_match = re.match(r"^##\s+(.+)$", line)
        if sub_match:
            if current_chapter is None:
                current_chapter, current_section = _ensure_default_section(chapters, current_chapter, current_section)
            current_section = SectionBlock(title=sub_match.group(1).strip())
            current_chapter.sections.append(current_section)
            active_special = None
            i += 1
            continue

        appendix_match = re.match(r'^\[APPENDIX\s+title="(.+?)"\]$', line)
        if appendix_match:
            appendix_title = appendix_match.group(1).strip()
            i += 1
            content_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "[/APPENDIX]":
                content_lines.append(lines[i])
                i += 1
            appendices.append((appendix_title, "\n".join(content_lines).strip()))
            i += 1
            continue

        if line.startswith("[FIGURE"):
            if current_section is None:
                current_chapter, current_section = _ensure_default_section(chapters, current_chapter, current_section)
            attrs = _parse_attrs(line)
            fig_title = attrs.get("title", "").strip() or "未命名圖片"
            fig_path = attrs.get("path", "").strip()
            if not fig_path:
                raise ParseError("圖片區塊缺少 path。")
            width_cm = float(attrs.get("width_cm", 14))
            i += 1
            note_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "[/FIGURE]":
                note_lines.append(lines[i])
                i += 1
            current_section.blocks.append(
                FigureBlock(
                    title=fig_title,
                    path=fig_path,
                    note="\n".join(note_lines).strip(),
                    width_cm=width_cm,
                )
            )
            i += 1
            continue

        if line.startswith("[TABLE"):
            if current_section is None:
                current_chapter, current_section = _ensure_default_section(chapters, current_chapter, current_section)
            attrs = _parse_attrs(line)
            table_title = attrs.get("title", "").strip() or "未命名表格"
            i += 1
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "[/TABLE]":
                table_lines.append(lines[i].rstrip())
                i += 1
            current_section.blocks.append(_parse_markdown_table(table_title, table_lines))
            i += 1
            continue

        if active_special == "abstract_zh":
            abstract_zh_lines.append(line)
        elif active_special == "abstract_en":
            abstract_en_lines.append(line)
        elif active_special == "acknowledgements":
            acknowledgements_lines.append(line)
        elif active_special == "references":
            references.append(line)
        elif active_special == "appendices":
            appendices.append((f"附錄{len(appendices) + 1}", line))
        else:
            if current_section is None:
                current_chapter, current_section = _ensure_default_section(chapters, current_chapter, current_section)
            current_section.blocks.append(line)
        i += 1

    if not chapters:
        current_chapter, current_section = _ensure_default_section(chapters, current_chapter, current_section)
        current_section.blocks.append("（原始內容空白）")

    return PaperContent(
        metadata=metadata,
        abstract_zh="\n".join(x for x in abstract_zh_lines if x).strip(),
        keywords_zh=keywords_zh,
        abstract_en="\n".join(x for x in abstract_en_lines if x).strip(),
        keywords_en=keywords_en,
        acknowledgements="\n".join(x for x in acknowledgements_lines if x).strip(),
        chapters=chapters,
        references=references,
        appendices=appendices,
    )


def _normalize_raw_text(raw_text: str) -> str:
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    normalized: list[str] = []
    saw_markdown_heading = any(line.strip().startswith("#") for line in lines)

    if saw_markdown_heading:
        return text

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            normalized.append("")
            continue

        meta_match = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if meta_match and meta_match.group(1).strip() in TITLE_KEYS:
            normalized.append(f"{meta_match.group(1).strip()}：{meta_match.group(2).strip()}")
            continue

        if line in SPECIAL_HEADINGS:
            normalized.append(f"# {line}")
            continue

        if re.match(r"^第[一二三四五六七八九十百0-9]+章", line):
            normalized.append(f"# {line}")
            continue

        if re.match(r"^第[一二三四五六七八九十百0-9]+節", line):
            normalized.append(f"## {line}")
            continue

        normalized.append(raw_line)

    converted = "\n".join(normalized)
    if re.search(r"^#\s+第[一二三四五六七八九十百0-9]+章", converted, re.MULTILINE):
        return converted

    meta_lines = [
        line
        for line in normalized
        if re.match(r"^([^:：]+)\s*：\s*(.+)$", line) and line.split("：", 1)[0] in TITLE_KEYS
    ]
    content_lines = [line for line in normalized if line not in meta_lines]

    wrapped: list[str] = []
    wrapped.extend(meta_lines)
    if wrapped:
        wrapped.append("")
    wrapped.append("# 第一章 內容整理")
    wrapped.append("## 第一節 主要內容")
    wrapped.extend(content_lines)
    return "\n".join(wrapped)


def _ensure_default_section(
    chapters: list[ChapterBlock],
    current_chapter: ChapterBlock | None,
    current_section: SectionBlock | None,
) -> tuple[ChapterBlock, SectionBlock]:
    if current_chapter is None:
        current_chapter = ChapterBlock(title="第一章 內容整理")
        chapters.append(current_chapter)
    if current_section is None:
        current_section = SectionBlock(title="第一節 主要內容")
        current_chapter.sections.append(current_section)
    return current_chapter, current_section


def _parse_attrs(line: str) -> dict[str, str]:
    return {k: v for k, v in re.findall(r'(\w+)="(.*?)"', line)}


def _parse_markdown_table(title: str, lines: Iterable[str]) -> TableBlock:
    filtered = [line.strip() for line in lines if line.strip()]
    note = ""
    grid_lines: list[str] = []
    for line in filtered:
        if line.startswith("註：") or line.startswith("Note."):
            note = line
        else:
            grid_lines.append(line)
    if len(grid_lines) < 2:
        raise ParseError(f"表格「{title}」資料不足。")
    headers = [cell.strip() for cell in grid_lines[0].strip("|").split("|")]
    rows: list[list[str]] = []
    for row_line in grid_lines[2:]:
        rows.append([cell.strip() for cell in row_line.strip("|").split("|")])
    return TableBlock(title=title, headers=headers, rows=rows, note=note)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    _configure_section(section)
    settings = document.settings.element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))


def _configure_section(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(2.5)
    section.footer_distance = Cm(2.5)


def _define_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "新細明體"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    normal.font.size = Pt(12)

    body = styles["BodyTextCustom"] if "BodyTextCustom" in styles else styles.add_style("BodyTextCustom", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = styles["Normal"]
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.first_line_indent = Cm(0.74)
    body.paragraph_format.space_after = Pt(0)
    body.font.name = "新細明體"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    body.font.size = Pt(12)

    heading1 = styles["Heading 1"]
    heading1.font.name = "新細明體"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.line_spacing = 1.5

    heading2 = styles["Heading 2"]
    heading2.font.name = "新細明體"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading2.paragraph_format.line_spacing = 1.5


def _build_cover(document: Document, paper: PaperContent, template: TemplateSpec) -> None:
    content = [
        paper.metadata.get("organization") or template.organization_default,
        paper.metadata.get("department") or "單位名稱",
        paper.metadata.get("document_type") or template.document_type,
        "",
        paper.metadata.get("title") or "文件題目",
        paper.metadata.get("title_en") or "Document Title",
        "",
        f"作者：{paper.metadata.get('author') or '未填寫'}",
        f"{paper.metadata.get('year') or '2026'} 年 {paper.metadata.get('month') or '6'} 月",
    ]
    for line in content:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(line)
        run.bold = True
        run.font.name = "新細明體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        run.font.size = Pt(18 if line in (paper.metadata.get("title"), paper.metadata.get("title_en")) else 16)
    document.add_page_break()


def _build_abstract_zh(document: Document, paper: PaperContent, template: TemplateSpec) -> None:
    _build_center_title(document, paper.metadata.get("title") or "文件題目")
    _build_center_subtitle(document, paper.metadata.get("author") or "未填寫")
    _build_center_subtitle(document, paper.metadata.get("department") or "單位名稱")
    _build_center_title(document, template.abstract_title, size=16)
    if paper.abstract_zh:
        _write_multiline_paragraphs(document, paper.abstract_zh)
        if paper.keywords_zh:
            p = document.add_paragraph(style="BodyTextCustom")
            run = p.add_run("關鍵詞：")
            run.bold = True
            p.add_run("、".join(paper.keywords_zh))
    else:
        p = document.add_paragraph(style="BodyTextCustom")
        p.add_run("（未提供摘要）")
    document.add_page_break()


def _build_abstract_en(document: Document, paper: PaperContent, template: TemplateSpec) -> None:
    _build_center_title(document, paper.metadata.get("title_en") or "Document Title")
    _build_center_subtitle(document, paper.metadata.get("author") or "Author")
    _build_center_subtitle(document, paper.metadata.get("department") or "Department")
    _build_center_title(document, template.english_abstract_title, size=16, east_asia_font="Times New Roman", latin_font="Times New Roman")
    if paper.abstract_en:
        _write_multiline_paragraphs(document, paper.abstract_en, font_name="Times New Roman")
        if paper.keywords_en:
            p = document.add_paragraph(style="BodyTextCustom")
            _set_run_font(p.add_run("Keywords: "), "Times New Roman", 12, italic=True)
            _set_run_font(p.add_run(", ".join(paper.keywords_en)), "Times New Roman", 12)
    else:
        p = document.add_paragraph(style="BodyTextCustom")
        _set_run_font(p.add_run("(English abstract not provided)"), "Times New Roman", 12)
    document.add_page_break()


def _build_center_title(
    document: Document,
    text: str,
    size: int = 24,
    east_asia_font: str = "新細明體",
    latin_font: str = "Times New Roman",
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = latin_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def _build_center_subtitle(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")


def _write_multiline_paragraphs(document: Document, text: str, font_name: str = "新細明體") -> None:
    for part in [x.strip() for x in text.split("\n") if x.strip()]:
        p = document.add_paragraph(style="BodyTextCustom")
        _set_run_font(p.add_run(part), font_name, 12)


def _insert_toc_field(paragraph, instruction: str) -> None:
    _append_field_run(paragraph, instruction, "請在 Word 中按 F9 更新目錄")


def _add_chapter_heading(document: Document, title: str, insert_page_break: bool) -> None:
    if insert_page_break:
        document.add_page_break()
    document.add_paragraph(title, style="Heading 1")


def _add_section_heading(document: Document, title: str) -> None:
    document.add_paragraph(title, style="Heading 2")


def _next_caption(prefix: str, chapter_no: int, counters: dict[int, int], title: str) -> str:
    counters[chapter_no] = counters.get(chapter_no, 0) + 1
    return f"{prefix} {chapter_no}-{counters[chapter_no]}  {title}"


def _add_caption(document: Document, caption: str, toc_id: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    _append_tc_field(p, caption, toc_id)
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")


def _add_table(document: Document, table_block: TableBlock) -> None:
    table = document.add_table(rows=1, cols=len(table_block.headers))
    table.style = "Table Grid"
    for idx, header in enumerate(table_block.headers):
        table.rows[0].cells[idx].text = header
    for row in table_block.rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(row_cells):
                row_cells[idx].text = value
    if table_block.note:
        p = document.add_paragraph(style="BodyTextCustom")
        p.add_run(table_block.note)
    document.add_paragraph("")


def _add_figure(document: Document, figure_block: FigureBlock) -> None:
    image_path = Path(figure_block.path)
    if not image_path.is_absolute():
        image_path = Path.cwd() / image_path
    if image_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Cm(figure_block.width_cm))
    else:
        p = document.add_paragraph(style="BodyTextCustom")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"[找不到圖片：{image_path}]")
    if figure_block.note:
        p = document.add_paragraph(style="BodyTextCustom")
        p.add_run(figure_block.note)
    document.add_paragraph("")


def _build_references(document: Document, references: list[str], title: str) -> None:
    if not references:
        return
    document.add_page_break()
    document.add_paragraph(title, style="Heading 1")
    for ref in references:
        p = document.add_paragraph(style="BodyTextCustom")
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.add_run(ref)


def _build_appendices(document: Document, appendices: list[tuple[str, str]], title: str) -> None:
    if not appendices:
        return
    document.add_page_break()
    document.add_paragraph(title, style="Heading 1")
    for appendix_title, appendix_content in appendices:
        document.add_paragraph(appendix_title, style="Heading 2")
        _write_multiline_paragraphs(document, appendix_content)


def _append_field_run(paragraph, instruction: str, display_text: str = "") -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    paragraph.add_run()._r.append(fld_begin)
    paragraph.add_run()._r.append(instr)
    paragraph.add_run()._r.append(fld_sep)
    if display_text:
        paragraph.add_run(display_text)
    paragraph.add_run()._r.append(fld_end)


def _append_tc_field(paragraph, text: str, toc_id: str) -> None:
    _append_field_run(paragraph, f'TC "{text}" \\f {toc_id} \\l 1')


def _set_page_number_format(section, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def _add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_field_run(p, "PAGE")


def _mark_fields_for_update(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _set_run_font(run, font_name: str, font_size: int, italic: bool = False) -> None:
    run.font.name = font_name
    east_asia_font = "Times New Roman" if font_name == "Times New Roman" else "新細明體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(font_size)
    run.italic = italic
